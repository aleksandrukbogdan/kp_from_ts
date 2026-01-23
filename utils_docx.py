
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def markdown_to_docx(md_text: str) -> io.BytesIO:
    """
    Converts Markdown text to a DOCX file in memory.
    """
    # 1. Convert Markdown to HTML
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')

    # 2. Create DOCX
    doc = Document()
    
    # Optional: Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 3. Traverse HTML and build DOCX
    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            if element.text.strip():
                doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'table':
            _add_table(doc, element)
            
    # 4. Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _add_table(doc, table_tag):
    """
    Helper to add a table from HTML tag to DOCX.
    """
    rows = table_tag.find_all('tr')
    if not rows:
        return

    # Determine columns count from first row
    cols_count = len(rows[0].find_all(['th', 'td']))
    if cols_count == 0:
        return

    docx_table = doc.add_table(rows=0, cols=cols_count)
    docx_table.style = 'Table Grid'

    for row in rows:
        cells = row.find_all(['th', 'td'])
        row_cells = docx_table.add_row().cells
        for i, cell in enumerate(cells):
            if i < len(row_cells):
                row_cells[i].text = cell.get_text(strip=True)
